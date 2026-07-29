"""Certificate editing: Certificates.docx -> Certificates good.pdf.

Automates SOP stage 5:

  * delete the first group of checkboxes
  * enter the current date under A.
  * under B., check Yes for (1)(2)(3)(4)(5)(9) and No for (6)(7)(8)(10)
  * save as "Certificates good.pdf"

Grouping assumption
-------------------
The SOP says "delete the first group of checkboxes" without saying how many
that is. Rather than hard-coding a count, the document's own structure is used:
checkboxes that appear **before the "A." paragraph** are the first group, and
the items under B. are numbered from the checkboxes that follow it.

That matches the document as described, but it is an inference. Run
``tegg inspect-docx <file>`` against a real certificate to see exactly what the
grouping resolves to before trusting a run, and set ``first_group_size`` in
config/workflow.yaml to pin it explicitly if the inference is wrong.
"""

from __future__ import annotations

import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

CHECKED_GLYPHS = {"☒", "☑"}
UNCHECKED_GLYPHS = {"☐"}


class CertificateError(Exception):
    """Raised when the certificate cannot be edited as the SOP describes."""


@dataclass
class Checkbox:
    """One checkbox and the paragraph it lives in."""

    element: object
    paragraph: object
    paragraph_index: int
    text: str
    kind: str  # "legacy" or "content-control"


@dataclass
class Structure:
    """What an inspection found in a certificate document."""

    document: object = None
    boxes: list[Checkbox] = field(default_factory=list)
    section_a_index: int | None = None
    section_b_index: int | None = None
    paragraphs: list = field(default_factory=list)

    @property
    def first_group(self) -> list[Checkbox]:
        """Checkboxes before section A."""
        if self.section_a_index is None:
            return []
        return [b for b in self.boxes if b.paragraph_index < self.section_a_index]

    @property
    def section_b_boxes(self) -> list[Checkbox]:
        """Checkboxes after section B, in document order."""
        if self.section_b_index is None:
            return [b for b in self.boxes if b not in self.first_group]
        return [b for b in self.boxes if b.paragraph_index > self.section_b_index]

    def describe(self) -> str:
        lines = [
            f"checkboxes found      : {len(self.boxes)}",
            f"kinds                 : {sorted({b.kind for b in self.boxes})}",
            f"section A paragraph   : {self.section_a_index}",
            f"section B paragraph   : {self.section_b_index}",
            f"first group (delete)  : {len(self.first_group)}",
            f"section B items       : {len(self.section_b_boxes)}",
            "",
            "first group:",
        ]
        lines += [f"  - {b.text[:70]}" for b in self.first_group] or ["  (none)"]
        lines += ["", "section B items (numbered in this order):"]
        lines += [
            f"  ({i}) {b.text[:66]}"
            for i, b in enumerate(self.section_b_boxes, 1)
        ] or ["  (none)"]
        return "\n".join(lines)


def inspect(source: Path) -> Structure:
    """Report the checkbox structure of a certificate without changing it."""
    source = Path(source)
    if not source.exists():
        raise CertificateError(f"certificate not found: {source}")

    document = Document(str(source))
    structure = Structure(document=document, paragraphs=list(document.paragraphs))

    for index, paragraph in enumerate(structure.paragraphs):
        stripped = paragraph.text.strip()
        if structure.section_a_index is None and stripped.startswith("A."):
            structure.section_a_index = index
        if structure.section_b_index is None and stripped.startswith("B."):
            structure.section_b_index = index

        for element in paragraph._p.findall(f".//{qn('w:checkBox')}"):
            structure.boxes.append(
                Checkbox(element, paragraph, index, stripped, "legacy")
            )
        for sdt in paragraph._p.findall(f".//{qn('w:sdt')}"):
            if sdt.findall(".//{*}checkbox"):
                structure.boxes.append(
                    Checkbox(sdt, paragraph, index, stripped, "content-control")
                )

    return structure


def set_checkbox(box: Checkbox, checked: bool) -> None:
    """Tick or untick one checkbox, whichever encoding it uses."""
    from docx.oxml import OxmlElement

    if box.kind == "legacy":
        for tag in ("w:checked", "w:default"):
            node = box.element.find(qn(tag))
            if node is not None:
                node.set(qn("w:val"), "1" if checked else "0")
                return
        node = OxmlElement("w:checked")
        node.set(qn("w:val"), "1" if checked else "0")
        box.element.append(node)
        return

    for node in box.element.findall(".//{*}checkbox"):
        for child in node.findall(".//{*}checked"):
            child.set(qn("w14:val"), "1" if checked else "0")
            return
    raise CertificateError("content-control checkbox has no <checked> element")


@dataclass
class CertificateEdit:
    """A record of what changed, for the run log and for verification."""

    date_written: str
    boxes_checked: list[int]
    boxes_unchecked: list[int]
    boxes_deleted: int

    def summary(self) -> str:
        return (
            f"date={self.date_written} yes={self.boxes_checked} "
            f"no={self.boxes_unchecked} deleted={self.boxes_deleted}"
        )


def edit_certificate(
    source: Path,
    output_docx: Path,
    *,
    yes_items: list[int],
    no_items: list[int],
    when: date | None = None,
    delete_first_group: bool = True,
    first_group_size: int | None = None,
) -> CertificateEdit:
    """Apply the SOP's certificate edits and save an edited .docx."""
    structure = inspect(source)
    if not structure.boxes:
        raise CertificateError(
            f"{Path(source).name}: no form checkboxes found. Run "
            "'tegg inspect-docx' on the file to see what it actually contains."
        )
    if structure.section_a_index is None:
        raise CertificateError(
            "could not find a paragraph beginning with 'A.' -- the date has "
            "nowhere to go, and the first checkbox group cannot be identified"
        )

    stamp = (when or date.today()).strftime("%m/%d/%Y")

    section_b = structure.section_b_boxes
    highest = max(yes_items + no_items)
    if len(section_b) < highest:
        raise CertificateError(
            f"section B has {len(section_b)} checkboxes but the SOP references "
            f"item ({highest}). Run 'tegg inspect-docx' to check the grouping."
        )

    for item in yes_items:
        set_checkbox(section_b[item - 1], True)
    for item in no_items:
        set_checkbox(section_b[item - 1], False)

    deleted = 0
    if delete_first_group:
        group = structure.first_group
        if first_group_size is not None:
            if len(group) != first_group_size:
                raise CertificateError(
                    f"expected a first group of {first_group_size} checkboxes but "
                    f"found {len(group)}. Run 'tegg inspect-docx' to inspect."
                )
            group = group[:first_group_size]
        for box in group:
            parent = box.paragraph._p.getparent()
            if parent is not None:
                parent.remove(box.paragraph._p)
                deleted += 1

    # The date goes into the A. paragraph, which survives the deletions above.
    _write_date(structure, stamp)

    output_docx = Path(output_docx)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    structure.document.save(str(output_docx))

    return CertificateEdit(
        date_written=stamp,
        boxes_checked=sorted(yes_items),
        boxes_unchecked=sorted(no_items),
        boxes_deleted=deleted,
    )


def _write_date(structure: Structure, stamp: str) -> None:
    if structure.section_a_index is None:
        raise CertificateError("no 'A.' paragraph to write the date into")
    structure.paragraphs[structure.section_a_index].add_run(f"  {stamp}")


def docx_to_pdf(source: Path, output_pdf: Path, timeout: int = 180) -> Path:
    """Convert a .docx to PDF using headless LibreOffice."""
    source = Path(source)
    output_pdf = Path(output_pdf)
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise CertificateError(
            "LibreOffice is required to convert the certificate to PDF but was "
            "not found on PATH. Install LibreOffice, or export the certificate "
            "to PDF by hand."
        )

    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory() as staging:
        # LibreOffice needs a writable user profile. Without an explicit one it
        # silently fails to load any document when HOME is unwritable or a
        # profile is already locked by another instance.
        profile = Path(staging) / "lo-profile"
        try:
            result = subprocess.run(
                [soffice, f"-env:UserInstallation={profile.as_uri()}",
                 "--headless", "--norestore", "--convert-to", "pdf",
                 "--outdir", staging, str(source)],
                capture_output=True, text=True, timeout=timeout, check=False,
            )
        except subprocess.TimeoutExpired as exc:
            raise CertificateError(
                f"LibreOffice timed out after {timeout}s converting {source.name}"
            ) from exc

        produced = Path(staging) / (source.stem + ".pdf")
        if not produced.exists():
            raise CertificateError(
                f"LibreOffice did not produce a PDF for {source.name}. "
                f"stderr: {result.stderr.strip()[:400]}"
            )
        shutil.move(str(produced), str(output_pdf))
    return output_pdf
