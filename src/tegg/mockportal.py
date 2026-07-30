"""A controlled stand-in for the live TEGG Pro portal.

This is a *test fixture that ships in the package*, because the autonomous mock
runner is a repository command and must not depend on anything outside it. It
never touches the network and serves only on 127.0.0.1.

Its job is to be wrong in the same ways the real portal is wrong, so that the
driver code exercised here is the same code that runs live. Every awkward
behaviour below was observed on the real site and is reproduced deliberately:

  * **Sign-in does not navigate.** Submitting posts to an auth endpoint from
    JavaScript and then re-renders in place. Any check made immediately after
    submit sees the form still there, which is what made real login look like
    rejected credentials for a whole day.
  * **The username field is ``<input type="email" name="email">``** whose
    accessible name comes from a wrapping tooltip label, and the visible text
    "User Name" is an unbound ``<label>``. ``get_by_label("User Name")`` finds
    nothing.
  * **The contractor ``<select>`` is ``required``.** Skip it and submit is
    silently refused -- indistinguishable from a bad password.
  * **A site must be selected first.** Until then every tab renders
    "PLEASE SEARCH FOR A CUSTOMER OR SITE USING THE SEARCH BAR".
  * **The search box is a typeahead driven by keystrokes.** Setting ``value``
    directly (Playwright's ``fill()``) raises no input events and produces no
    results, so the driver must type.
  * **Report leaves are ``<a>`` tags with no ``href``**, inside
    ``tr.child-level-1``, *expanded by default*. Clicking the parent collapses
    them. ``get_by_role("link")`` will not match them.
  * **Print Report opens a popup after a delay** and does not download.
  * **The SSRS toolbar is not usable immediately**; the format ``<select>``
    stays disabled while the viewer renders.
  * **Export delivers the PDF as an inline response**, with no download event
    and no navigation, so only a ``response`` listener sees it.

Faults can be injected per report so the failure paths are exercised
deterministically rather than waited for. See :class:`Faults`.
"""

from __future__ import annotations

import argparse
import io
import json
import threading
from dataclasses import dataclass, field
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from urllib.parse import parse_qs, urlparse

from . import canonical

# --- the account the mock accepts ------------------------------------------
# Obviously fake, and never read from the environment. A test that needs a
# wrong password passes its own string; nothing here is a real credential.
USERNAME = "mock-technician"
PASSWORD = "mock-password-not-a-real-secret"
CONTRACTOR_LABEL = "TEGGPro Lippolis"
CONTRACTOR_VALUE = "TEGGProLippolis595"

CUSTOMER = "Northwind Manufacturing"
SITE = "Bay Street Plant"
AGREEMENT = "STD00000000XX-01/26-01"
SITE_VISIT = "T26-001"
VISIT_DATE = "2026-03-11"
CONTRACTOR_ADDRESS = "Example Electric, Inc., 1 Test Way, Springfield"

# Other sites, so selecting the right one is a real choice rather than the only
# option. "Bay Street Plant" is deliberately not discoverable by typing the
# customer name, mirroring the live portal.
OTHER_SITES = ["Riverside Depot", "Harbour Works"]
OTHER_CUSTOMERS = ["Northwind Manufacturing", "Southgate Logistics"]


@dataclass(frozen=True)
class MockReport:
    """One entry in the Standard ESA Reports list."""

    key: str
    label: str
    slug: str
    pages: int
    # Report leaves nested under an accordion parent, e.g. Short/Long Form.
    parent: str = ""
    # Which parameter dropdowns this report's form shows.
    fields: tuple[str, ...] = ("Agreement",)
    # Whether the generated PDF carries the site-visit identifier. The real
    # Short Form does not -- its header is "<site> - <agreement>" only, which
    # was confirmed by extracting text from the live 70-page export.
    carries_visit_id: bool = True


REPORTS: tuple[MockReport, ...] = (
    MockReport(
        key=canonical.PROBLEM_COUNT_SUMMARY,
        label="Problem Count Summary",
        slug="pcs",
        pages=2,
        fields=("Agreement",),
    ),
    MockReport(
        key=canonical.EQUIPMENT_INVENTORY_SHORT,
        label="Short Form",
        slug="eis",
        pages=5,
        parent="Equipment Inventory",
        fields=("Agreement", "Order By", "Images"),
        carries_visit_id=False,
    ),
    MockReport(
        key=canonical.EQUIPMENT_INVENTORY_LONG,
        label="Long Form",
        slug="eil",
        pages=8,
        parent="Equipment Inventory",
        fields=("Agreement", "Order By", "Images"),
        carries_visit_id=False,
    ),
    MockReport(
        key=canonical.STANDARD_IR,
        label="Standard IR Report",
        slug="sir",
        # At least two pages, or the cover cannot be split off.
        pages=6,
        fields=("Agreement",),
    ),
    MockReport(
        key=canonical.EQUIPMENT_ITEM_PROBLEMS,
        label="Equipment Item Problems",
        slug="eip",
        pages=4,
        fields=("Agreement", "Images"),
    ),
    MockReport(
        key=canonical.EDS_ALL_PROBLEMS,
        label="EDS Component Problem Summary",
        slug="eds",
        pages=3,
        fields=("Agreement",),
    ),
)

BY_SLUG = {r.slug: r for r in REPORTS}
BY_KEY = {r.key: r for r in REPORTS}

FIELD_OPTIONS = {
    "Agreement": [AGREEMENT, "STD00000000XX-01/25-04"],
    # "Locations" plus "Tag ID" is how the Order By control is recognised.
    "Order By": ["Locations", "Tag ID", "Equipment Type"],
    "Images": ["Include Images", "No Images"],
}

# --- fault names, so a caller cannot invent one silently -------------------
NO_VIEWER = "no_viewer"
NO_PDF = "no_pdf"
HTML_BODY = "html_body"
EMPTY_PDF = "empty_pdf"
MISSING_AGREEMENT = "missing_agreement"
REPORT_ABSENT = "report_absent"
MALFORMED_INPUT = "malformed_input"

FAULT_NAMES = frozenset(
    {
        NO_VIEWER,
        NO_PDF,
        HTML_BODY,
        EMPTY_PDF,
        MISSING_AGREEMENT,
        REPORT_ABSENT,
        MALFORMED_INPUT,
    }
)

# What a report renders when the technician's recommendation and estimate came
# through unresolved. This is a real shape, not an invented one: the live
# reports interpolate the recommendation text and the dollar estimate, so a row
# whose source value was blank or non-numeric surfaces as the unsubstituted
# token rather than as an error. It is a *valid* PDF, which is the point --
# only reading the text catches it.
MALFORMED_RECOMMENDATION = "Recommendation: {{recommendation}}"
MALFORMED_ESTIMATE = "Repair estimate: $TBD_ (TODO: pricing not returned)"


@dataclass
class Faults:
    """Deterministic failure injection, per report slug."""

    per_report: dict[str, str] = field(default_factory=dict)

    def __post_init__(self) -> None:
        unknown = {v for v in self.per_report.values()} - FAULT_NAMES
        if unknown:
            raise ValueError(
                f"unknown fault(s) {sorted(unknown)}; choose from "
                f"{sorted(FAULT_NAMES)}"
            )
        bad = set(self.per_report) - set(BY_SLUG)
        if bad:
            raise ValueError(f"unknown report slug(s) {sorted(bad)}")

    def of(self, slug: str) -> str:
        return self.per_report.get(slug, "")


@dataclass
class Timing:
    """How long the portal's slow steps take, in milliseconds.

    The live values are in the comments. Tests use small ones; the point is that
    the *shape* is reproduced -- a popup that appears late and a toolbar that is
    not ready when it appears -- not the exact duration.
    """

    login_ms: int = 120          # live: ~750
    print_report_ms: int = 400   # live: ~20_000
    viewer_ready_ms: int = 350    # live: several seconds after networkidle
    list_load_ms: int = 150       # live: the async site-visit table


# ---------------------------------------------------------------------------
# PDF and certificate generation
# ---------------------------------------------------------------------------


def make_pdf(
    report: MockReport, pages: int | None = None, *, malformed: bool = False
) -> bytes:
    """A PDF carrying the fields validation will look for.

    Which fields appear mirrors the live export: the header is
    ``<site> - <agreement>`` plus the contractor's own address. The customer
    name is *not* on the Equipment Inventory forms, and neither is the visit id,
    which is why validation checks the site and the agreement rather than the
    customer.
    """
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas as pdfcanvas

    total = pages if pages is not None else report.pages
    buffer = io.BytesIO()
    pdf = pdfcanvas.Canvas(buffer, pagesize=LETTER)
    width, height = LETTER
    for page in range(1, total + 1):
        pdf.setFont("Helvetica-Bold", 15)
        pdf.drawCentredString(width / 2, height - 90, canonical.title(report.key))
        pdf.setFont("Helvetica", 11)
        lines = [
            f"{SITE} - {AGREEMENT}",
            CONTRACTOR_ADDRESS,
        ]
        if report.carries_visit_id:
            lines.append(f"Site Visit: {SITE_VISIT}   Visit date: {VISIT_DATE}")
        lines.append(f"Page {page} of {total}")
        if malformed:
            lines += [MALFORMED_RECOMMENDATION, MALFORMED_ESTIMATE]
        offset = height - 130
        for line in lines:
            pdf.drawCentredString(width / 2, offset, line)
            offset -= 20
        pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def make_certificate_docx() -> bytes:
    """A certificate whose section B boxes are Wingdings glyphs, two per item.

    Eleven items, two glyphs each, exactly as the real document is built. This
    is the shape that cannot be ticked programmatically without a proven
    mapping, and reproducing it here is what keeps the honest blocker honest.
    """
    from docx import Document

    document = Document()
    document.add_heading("TEGG Certificate of Inspection", level=1)
    document.add_paragraph(f"Customer: {CUSTOMER}")
    document.add_paragraph(f"Site: {SITE}")
    document.add_paragraph(f"Agreement: {AGREEMENT}")
    document.add_paragraph(f"Site Visit: {SITE_VISIT}")
    document.add_paragraph("A.  Date of inspection: ____________________")
    document.add_paragraph("B.  Certification items")
    for index in range(1, 12):
        paragraph = document.add_paragraph(f"({index}) certification item {index}  ")
        for _ in range(2):
            run = paragraph.add_run(" ")  # Wingdings empty box
            run.font.name = "Wingdings"
    document.add_paragraph("Authorised signature: ____________________")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Page templates
# ---------------------------------------------------------------------------

LOGIN_HTML = """<!doctype html>
<html><head><meta charset="utf-8"><title>TEGGPro 2.0</title></head><body>
<div id="app">
<form id="signin" class="smart-form client-form" novalidate>
  <header>Sign in to TEGGPro</header>
  <fieldset>
    <section>
      <!-- Unbound label: get_by_label("User Name") will not find the input. -->
      <label class="label">User Name</label>
      <label class="input">
        <input autofocus name="email" type="email" data-required>
        <b class="tooltip tooltip-top-right">Please enter email address/username</b>
      </label>
    </section>
    <section>
      <label class="label">Password</label>
      <label class="input">
        <input name="password" type="password" data-required>
        <b class="tooltip tooltip-top-right">Enter your password</b>
      </label>
    </section>
    <section>
      <label class="label">Contractor</label>
      <select class="form-control" name="contractorConnection" required
              type="contractorConnection">
        <option disabled value="none">Select Contractor</option>
        __CONTRACTORS__
      </select>
    </section>
  </fieldset>
  <footer>
    <div id="login-error" class="alert alert-danger" style="display:none"></div>
    <button class="btn btn-primary" name="btnLogin">Sign in</button>
  </footer>
</form>
</div>
<script>
// The real page authenticates with fetch() and re-renders in place. There is no
// navigation, so a driver must poll for an outcome.
var form = document.getElementById('signin');
var errorBox = document.getElementById('login-error');
form.addEventListener('submit', function (event) {
  event.preventDefault();
  errorBox.style.display = 'none';
  var contractor = form.contractorConnection.value;
  // Required and silently enforced: no request is even made without it.
  if (!contractor || contractor === 'none') { return; }
  fetch('/api/1.0/technician/auth-new', {
    method: 'POST',
    headers: {'Content-Type': 'application/json'},
    body: JSON.stringify({
      email: form.email.value,
      password: form.password.value,
      contractorConnection: contractor
    })
  }).then(function (r) { return r.json(); }).then(function (data) {
    if (!data.ok) {
      errorBox.textContent = data.message || 'Invalid credentials.';
      errorBox.style.display = 'block';
      return;
    }
    setTimeout(function () {
      // The form goes away and the route changes -- both signals a driver can
      // legitimately wait for.
      document.getElementById('app').innerHTML =
        '<h1>Dashboard</h1><p>Signed in.</p>';
      history.pushState({}, '', '/sales/dashboard');
    }, __LOGIN_MS__);
  });
});
</script>
</body></html>
"""


DOCUMENTATION_HTML = """<!doctype html>
<html><head><meta charset="utf-8"><title>Documentation - TEGGPro</title></head>
<body>
<nav><a href="/sales/documentation">Documentation</a></nav>
<div class="search-bar">
  <input id="site-search" type="text" autocomplete="off"
         placeholder="Enter customer or site name">
  <typeahead-container id="typeahead" style="display:none"></typeahead-container>
</div>
<div id="selected"></div>
<div class="tabs">
  <span class="tab" id="tab-reports">Reports</span>
  <span class="tab" id="tab-library">Document Library</span>
</div>
<div id="panel"><p>PLEASE SEARCH FOR A CUSTOMER OR SITE USING THE SEARCH BAR</p></div>
<iframe id="sink" style="display:none"></iframe>
<script>
var DATA = __DATA__;
var state = {site: null, tab: null, report: null, collapsed: false};
var box = document.getElementById('site-search');
var container = document.getElementById('typeahead');
var panel = document.getElementById('panel');

function unselected() {
  return '<p>PLEASE SEARCH FOR A CUSTOMER OR SITE USING THE SEARCH BAR</p>';
}

// ngx-bootstrap typeahead: driven by input events, so assigning .value does
// nothing. Results are grouped Customers / Sites.
box.addEventListener('input', function () {
  var q = box.value.trim().toLowerCase();
  if (q.length < 2) { container.style.display = 'none'; return; }
  var customers = DATA.customers.filter(function (c) {
    return c.toLowerCase().indexOf(q) !== -1;
  });
  // Searching the customer name does NOT surface its sites, exactly as live.
  var sites = DATA.sites.filter(function (s) {
    return s.toLowerCase().indexOf(q) !== -1;
  });
  var html = '';
  if (customers.length) {
    html += '<h6 class="dropdown-header">Customers</h6>';
    customers.forEach(function (c) {
      html += '<a class="dropdown-item" data-kind="customer">' + c + '</a>';
    });
  }
  if (sites.length) {
    html += '<h6 class="dropdown-header">Sites</h6>';
    sites.forEach(function (s) {
      html += '<a class="dropdown-item" data-kind="site">' + s + '</a>';
    });
  }
  container.innerHTML = html;
  container.style.display = html ? 'block' : 'none';
});

container.addEventListener('click', function (event) {
  var item = event.target.closest('.dropdown-item');
  if (!item) { return; }
  container.style.display = 'none';
  if (item.dataset.kind !== 'site') { return; }
  state.site = item.textContent.trim();
  document.getElementById('selected').textContent = 'Site: ' + state.site;
  panel.innerHTML = '<p>Select a tab.</p>';
});

document.getElementById('tab-reports').addEventListener('click', function () {
  if (!state.site) { panel.innerHTML = unselected(); return; }
  state.tab = 'reports';
  panel.innerHTML = '<p class="tab-link" id="standard-esa">Standard ESA Reports</p>';
  document.getElementById('standard-esa').addEventListener('click', renderList);
});

document.getElementById('tab-library').addEventListener('click', function () {
  if (!state.site) { panel.innerHTML = unselected(); return; }
  state.tab = 'library';
  panel.innerHTML = '<p><a id="cert-link" href="/download/certificate">Certificates</a></p>';
});

function renderList() {
  state.collapsed = false;
  var rows = '';
  var parents = {};
  DATA.reports.forEach(function (r) {
    if (r.absent) { return; }
    if (r.parent) {
      if (!parents[r.parent]) {
        parents[r.parent] = true;
        rows += '<tr class="parent-level"><td><span class="parent" ' +
                'data-parent="' + r.parent + '">' + r.parent + '</span></td></tr>';
      }
      // No href, so get_by_role("link") cannot see it. Expanded by default.
      rows += '<tr class="child-level-1" data-parent="' + r.parent + '">' +
              '<td><a data-slug="' + r.slug + '">' + r.label + '</a></td></tr>';
    } else {
      rows += '<tr class="parent-level"><td><a data-slug="' + r.slug + '">' +
              r.label + '</a></td></tr>';
    }
  });
  panel.innerHTML = '<table id="report-list"><tbody>' + rows + '</tbody></table>';

  // Clicking the accordion parent COLLAPSES the children that were already
  // visible. This is what made early attempts time out.
  panel.querySelectorAll('.parent').forEach(function (p) {
    p.addEventListener('click', function () {
      state.collapsed = !state.collapsed;
      panel.querySelectorAll('tr.child-level-1').forEach(function (row) {
        row.style.display = state.collapsed ? 'none' : '';
      });
    });
  });
  panel.querySelectorAll('a[data-slug]').forEach(function (a) {
    a.addEventListener('click', function () { openReport(a.dataset.slug); });
  });
}

function openReport(slug) {
  var report = DATA.reports.filter(function (r) { return r.slug === slug; })[0];
  state.report = report;
  var html = '<h2 id="report-title">' + report.label + '</h2>';
  report.fields.forEach(function (f) {
    html += '<p><select data-field="' + f + '"><option value=""></option>';
    (DATA.options[f] || []).forEach(function (o) {
      html += '<option value="' + o + '">' + o + '</option>';
    });
    html += '</select></p>';
  });
  html += '<p><span id="print-report">Print Report</span></p>';
  html += '<p id="print-error" style="display:none"></p>';
  panel.innerHTML = html;
  document.getElementById('print-report').addEventListener('click', function () {
    var unset = [];
    panel.querySelectorAll('select[data-field]').forEach(function (s) {
      if (!s.value) { unset.push(s.dataset.field); }
    });
    if (unset.length) {
      // Silently refuses, like the real form. No error is shown.
      return;
    }
    if (report.fault === 'no_viewer') { return; }
    setTimeout(function () {
      window.open('/Report/SSOReportViewer.aspx?db=MockDb&rt=rv&report=' +
                  report.slug, '_blank');
    }, DATA.timing.print_report_ms);
  });
}
</script>
</body></html>
"""


VIEWER_HTML = """<!doctype html>
<html><head><meta charset="utf-8"><title>TEGGPro Reports</title></head><body>
<div id="ReportViewer1">
  <div id="ReportViewer1_ctl01">
    <div id="ReportViewer1_ctl01_ctl05">
      <select id="ReportViewer1_ctl01_ctl05_ctl00" title="Export Formats" disabled>
        <option value="">Select a format</option>
        <option value="PDF">Acrobat (PDF) file</option>
        <option value="EXCELOPENXML">Excel</option>
        <option value="WORDOPENXML">Word</option>
      </select>
      <a id="ReportViewer1_ctl01_ctl05_ctl01" title="Export">Export</a>
    </div>
  </div>
  <div id="report-body">Loading report...</div>
</div>
<iframe id="export-sink" style="display:none"></iframe>
<script>
var SLUG = __SLUG__;
var READY_MS = __READY_MS__;
var format = document.getElementById('ReportViewer1_ctl01_ctl05_ctl00');
var link = document.getElementById('ReportViewer1_ctl01_ctl05_ctl01');

// The toolbar is present but unusable while the viewer renders, so a driver
// that selects a format immediately after networkidle fails.
setTimeout(function () {
  format.disabled = false;
  document.getElementById('report-body').textContent = 'Report rendered.';
}, READY_MS);

link.addEventListener('click', function (event) {
  event.preventDefault();
  // Order matters on the real toolbar: with no format chosen, Export does
  // nothing at all.
  if (format.value !== 'PDF') { return; }
  // Delivered as an inline response into a hidden iframe: no download event,
  // no navigation. Only a response listener sees this.
  document.getElementById('export-sink').src =
    '/Reserved.ReportViewerWebControl.axd?ReportSession=mock' +
    '&ControlID=mock&OpType=Export&Format=PDF&report=' + SLUG;
});
</script>
</body></html>
"""


def _json(value: object) -> str:
    return json.dumps(value)


class Handler(BaseHTTPRequestHandler):
    protocol_version = "HTTP/1.1"

    def log_message(self, *args) -> None:  # keep test output clean
        pass

    # -- helpers ----------------------------------------------------------
    @property
    def options(self):
        return self.server.options

    def _send(
        self,
        body: bytes,
        content_type: str = "text/html; charset=utf-8",
        extra: dict | None = None,
        status: int = 200,
    ) -> None:
        self.send_response(status)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(len(body)))
        for key, value in (extra or {}).items():
            self.send_header(key, value)
        self.end_headers()
        self.wfile.write(body)

    def _authed(self) -> bool:
        return "tegg-mock-session=1" in self.headers.get("Cookie", "")

    # -- routing ----------------------------------------------------------
    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        path = parsed.path
        query = parse_qs(parsed.query)

        if path in ("/", "/auth/login"):
            return self._send(self._login_page())

        if not self._authed():
            # An unauthenticated SPA route renders the sign-in page rather than
            # redirecting, which is what the real app does.
            return self._send(self._login_page())

        if path in ("/sales/dashboard", "/sales/gm-dashboard"):
            return self._send(b"<h1>Dashboard</h1><p>Signed in.</p>")

        if path == "/sales/documentation":
            return self._send(self._documentation_page())

        if path == "/reports":
            # The dead end the original code assumed was the route.
            return self._send(
                b"<h1>Reports</h1><p>There are currently no agreements to "
                b"report on.</p>"
            )

        if path == "/Report/SSOReportViewer.aspx":
            slug = (query.get("report") or [""])[0]
            if slug not in BY_SLUG:
                return self._send(b"<p>No such report.</p>", status=404)
            return self._send(self._viewer_page(slug))

        if path == "/Reserved.ReportViewerWebControl.axd":
            return self._export(query)

        if path == "/download/certificate":
            return self._send(
                self.server.certificate_bytes,
                self.server.certificate_content_type,
                {
                    "Content-Disposition":
                        f'attachment; filename="{self.server.certificate_name}"'
                },
            )

        return self._send(b"<p>Nothing here.</p>", status=404)

    def do_POST(self) -> None:
        path = urlparse(self.path).path
        length = int(self.headers.get("Content-Length", 0) or 0)
        raw = self.rfile.read(length) if length else b""

        if path == "/api/1.0/technician/auth-new":
            try:
                payload = json.loads(raw or b"{}")
            except ValueError:
                payload = {}
            ok = (
                payload.get("email") == self.options.username
                and payload.get("password") == self.options.password
                and payload.get("contractorConnection") == CONTRACTOR_VALUE
            )
            body = _json(
                {"ok": ok, "message": "" if ok else "Invalid credentials."}
            ).encode()
            extra = {"Set-Cookie": "tegg-mock-session=1; Path=/"} if ok else {}
            return self._send(body, "application/json", extra)

        return self._send(b"<p>Nothing here.</p>", status=404)

    # -- the export endpoint ---------------------------------------------
    def _export(self, query: dict) -> None:
        slug = (query.get("report") or [""])[0]
        report = BY_SLUG.get(slug)
        if report is None:
            return self._send(b"<p>No such report.</p>", status=404)

        fault = self.options.faults.of(slug)
        if fault == NO_PDF:
            # Generation "fails" server-side: a 500 with no body, which is what
            # a report that never arrives actually looks like.
            return self._send(b"", "text/plain", status=500)
        if fault == HTML_BODY:
            # The single most common real failure: an error page served with a
            # PDF content type, which only a magic-byte check catches.
            return self._send(
                b"<html><body><h1>Server Error</h1>"
                b"<p>Report generation failed.</p></body></html>",
                "application/pdf",
            )
        if fault == EMPTY_PDF:
            return self._send(b"", "application/pdf")

        # A well-formed PDF carrying unresolved recommendation/estimate tokens.
        # Export succeeds; only validation's text pass can catch this one.
        return self._send(
            make_pdf(report, malformed=(fault == MALFORMED_INPUT)), "application/pdf"
        )

    # -- pages ------------------------------------------------------------
    def _login_page(self) -> bytes:
        contractors = "".join(
            f'<option value="{value}">{label}</option>'
            for value, label in (
                ("TEGGProADI396", "TEGGPro ADI"),
                (CONTRACTOR_VALUE, CONTRACTOR_LABEL),
                ("TEGGProMatco277", "TEGGPro Matco"),
            )
        )
        return (
            LOGIN_HTML.replace("__CONTRACTORS__", contractors)
            .replace("__LOGIN_MS__", str(self.options.timing.login_ms))
            .encode()
        )

    def _documentation_page(self) -> bytes:
        reports = []
        for report in REPORTS:
            fault = self.options.faults.of(report.slug)
            fields = list(report.fields)
            options = dict(FIELD_OPTIONS)
            reports.append(
                {
                    "slug": report.slug,
                    "label": report.label,
                    "parent": report.parent,
                    "fields": fields,
                    "fault": fault,
                    "absent": fault == REPORT_ABSENT,
                }
            )

        options = {name: list(values) for name, values in FIELD_OPTIONS.items()}
        # A report whose agreement option is missing is how "the wrong site is
        # selected" presents on the real portal.
        if MISSING_AGREEMENT in self.options.faults.per_report.values():
            options["Agreement"] = [
                o for o in options["Agreement"] if o != AGREEMENT
            ] or ["STD00000000XX-01/25-04"]

        data = {
            "customers": OTHER_CUSTOMERS,
            "sites": [SITE, *OTHER_SITES],
            "reports": reports,
            "options": options,
            "timing": {"print_report_ms": self.options.timing.print_report_ms},
        }
        return DOCUMENTATION_HTML.replace("__DATA__", _json(data)).encode()

    def _viewer_page(self, slug: str) -> bytes:
        return (
            VIEWER_HTML.replace("__SLUG__", _json(slug))
            .replace("__READY_MS__", str(self.options.timing.viewer_ready_ms))
            .encode()
        )


@dataclass
class Options:
    """Everything about this mock a test may want to vary."""

    username: str = USERNAME
    password: str = PASSWORD
    faults: Faults = field(default_factory=Faults)
    timing: Timing = field(default_factory=Timing)


class MockPortal:
    """Context manager running the mock portal on a free localhost port."""

    def __init__(
        self,
        port: int = 0,
        *,
        options: Options | None = None,
        certificate_bytes: bytes | None = None,
        certificate_name: str = "Certificates good.docx",
        certificate_content_type: str = (
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
    ) -> None:
        self.server = ThreadingHTTPServer(("127.0.0.1", port), Handler)
        self.server.options = options or Options()
        self.server.certificate_bytes = (
            certificate_bytes
            if certificate_bytes is not None
            else make_certificate_docx()
        )
        self.server.certificate_name = certificate_name
        self.server.certificate_content_type = certificate_content_type
        self.port = self.server.server_address[1]
        self.url = f"http://127.0.0.1:{self.port}"

    @property
    def options(self) -> Options:
        return self.server.options

    def __enter__(self) -> "MockPortal":
        self.thread = threading.Thread(target=self.server.serve_forever, daemon=True)
        self.thread.start()
        return self

    def __exit__(self, *exc) -> None:
        self.server.shutdown()
        self.server.server_close()


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("--port", type=int, default=8899)
    args = parser.parse_args()
    with MockPortal(args.port) as portal:
        print(f"mock TEGG portal on {portal.url}  (ctrl-c to stop)")
        print(f"  sign in as {USERNAME} / <see tegg.mockportal.PASSWORD>")
        try:
            threading.Event().wait()
        except KeyboardInterrupt:
            pass
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
