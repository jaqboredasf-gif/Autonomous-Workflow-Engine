"""A stand-in TEGG Pro portal for testing the automation driver.

This reproduces the navigation, dropdowns and export flow that Paul's SOP
describes, so `tegg.portal` can be exercised end to end without touching the
real site. It is a test fixture, not a simulation of TEGG Pro -- the labels and
control types come straight from the SOP, which is exactly what the driver
keys off.

Run standalone for manual poking:

    python tests/mock_portal.py --port 8899
"""

from __future__ import annotations

import argparse
import io
import threading
from http.server import BaseHTTPRequestHandler, HTTPServer
from urllib.parse import parse_qs, urlparse

from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

USERNAME = "PLippolis1"
PASSWORD = "test-password"
CONTRACTOR = "Lippolis"

AGREEMENTS = ["AG-118422", "AG-220015"]
SITE_VISITS = ["2026-05-14", "2026-02-03"]
CONTACTS = ["Dana Reyes", "Sam Cole"]
SITES = ["Plant 3 - Toledo", "Plant 7 - Akron"]

# report label -> (url slug, download filename, page count)
REPORTS = {
    "Equipment Inventory and Short Form": ("eis", "EquipmentInventoryShortForm.pdf", 5),
    "Equipment Inventory and Long Form": ("eil", "EquipmentInventoryLongForm.pdf", 12),
    "Equipment Item Problems and Include All Images": (
        "eip", "EquipmentItemProblem_AllImages.pdf", 7,
    ),
    "Problem Count Summary": ("pcs", "ProblemCountSummary.pdf", 2),
    "Standard IR Report": ("sir", "StandardIRReport.pdf", 9),
    "EDS Component Problem Summary and All Problems": (
        "eds", "EDSAllProblemsOnly.pdf", 4,
    ),
}
SLUG_TO_REPORT = {slug: name for name, (slug, _, _) in REPORTS.items()}

# Which dropdowns each report shows, mirroring the SOP's per-report parameters.
REPORT_FIELDS = {
    "Equipment Inventory and Short Form": ["Agreement", "Site Visit", "Order By", "Images"],
    "Equipment Inventory and Long Form": ["Agreement", "Record Selection"],
    "Equipment Item Problems and Include All Images": ["Contact", "Agreement", "Site Visit"],
    "Problem Count Summary": ["Contact", "Agreement", "Site Visit"],
    "Standard IR Report": ["Contact", "Agreement", "Site Visit"],
    "EDS Component Problem Summary and All Problems": ["Contact", "Agreement", "Site Visit"],
}

FIELD_OPTIONS = {
    "Agreement": AGREEMENTS,
    "Site Visit": SITE_VISITS,
    "Contact": CONTACTS,
    "Order By": ["Locations", "Equipment Type"],
    "Images": ["Include Images", "No Images"],
    "Record Selection": ["Site Visits", "All Records"],
    "Solution": AGREEMENTS,
}


def make_pdf(title: str, pages: int) -> bytes:
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    width, height = LETTER
    for page in range(1, pages + 1):
        pdf.setFont("Helvetica-Bold", 20)
        pdf.drawCentredString(width / 2, height - 180, title)
        pdf.setFont("Helvetica", 12)
        pdf.drawCentredString(width / 2, height - 210, f"page {page} of {pages}")
        pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def make_docx() -> bytes:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document = Document()
    document.add_paragraph("TEGG Certificate of Inspection")

    # First checkbox group -- the one the SOP says to delete.
    document.add_paragraph("Preliminary items")
    for index in range(3):
        _add_checkbox_paragraph(document, f"preliminary item {index + 1}", OxmlElement, qn)

    document.add_paragraph("A.  Date of inspection:")
    document.add_paragraph("B.  Certification items")
    for index in range(10):
        _add_checkbox_paragraph(document, f"({index + 1}) certification item", OxmlElement, qn)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _add_checkbox_paragraph(document, text, OxmlElement, qn):
    """Append a paragraph carrying a legacy <w:checkBox> form field."""
    paragraph = document.add_paragraph(text)
    run = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    ffdata = OxmlElement("w:ffData")
    checkbox = OxmlElement("w:checkBox")
    default = OxmlElement("w:default")
    default.set(qn("w:val"), "0")
    checkbox.append(default)
    ffdata.append(checkbox)
    fld.append(ffdata)
    run._r.append(fld)


def _page(title: str, body: str) -> bytes:
    return f"""<!doctype html><html><head><meta charset="utf-8">
<title>{title} - TEGG Pro</title></head><body>
<nav><a href="/sales/gm-dashboard">Dashboard</a>
<a href="/documentation">Documentation</a>
<a href="/document-library">Document Library</a>
<a href="/reports">Reports</a></nav>
<h1>{title}</h1>
{body}
</body></html>""".encode()


def _select(label: str, options: list[str]) -> str:
    choices = "".join(f'<option value="{o}">{o}</option>' for o in options)
    field_id = label.lower().replace(" ", "-")
    return (
        f'<p><label for="{field_id}">{label}</label>'
        f'<select id="{field_id}" name="{field_id}">'
        f'<option value=""></option>{choices}</select></p>'
    )


class Handler(BaseHTTPRequestHandler):
    def log_message(self, *args):  # keep test output quiet
        pass

    # -- helpers ----------------------------------------------------------
    def _send(self, body: bytes, content_type="text/html; charset=utf-8", extra=None):
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(len(body)))
        for key, value in (extra or {}).items():
            self.send_header(key, value)
        self.end_headers()
        self.wfile.write(body)

    def _redirect(self, location):
        self.send_response(302)
        self.send_header("Location", location)
        self.end_headers()

    def _authed(self) -> bool:
        return "tegg-session=1" in self.headers.get("Cookie", "")

    def _form(self) -> dict:
        length = int(self.headers.get("Content-Length", 0))
        raw = self.rfile.read(length).decode()
        return {k: v[0] for k, v in parse_qs(raw).items()}

    # -- routing ----------------------------------------------------------
    def do_GET(self):
        path = urlparse(self.path).path
        query = parse_qs(urlparse(self.path).query)

        if path in ("/", "/auth/login"):
            return self._send(self._login_page())

        if not self._authed():
            return self._redirect("/auth/login")

        if path == "/sales/gm-dashboard":
            return self._send(_page("GM Dashboard", "<p>Welcome.</p>"))

        if path == "/documentation":
            return self._send(self._documentation_page())

        if path == "/document-library":
            return self._send(
                _page("Document Library", '<a href="/document-library/certificates">Certificates</a>')
            )

        if path == "/document-library/certificates":
            return self._send(self._certificates_page())

        if path == "/reports":
            return self._send(
                _page("Reports", '<a href="/reports/standard-esa">Standard ESA Reports</a>')
            )

        if path == "/reports/standard-esa":
            links = "".join(
                f'<p><a href="/reports/standard-esa/{slug}">{name}</a></p>'
                for name, (slug, _, _) in REPORTS.items()
            )
            return self._send(_page("Standard ESA Reports", links))

        if path.startswith("/reports/standard-esa/"):
            slug = path.rsplit("/", 1)[-1]
            if slug in SLUG_TO_REPORT:
                staged = query.get("staged", ["0"])[0] == "1"
                return self._send(self._report_page(SLUG_TO_REPORT[slug], staged))

        self.send_error(404)

    def do_POST(self):
        path = urlparse(self.path).path
        form = self._form()

        if path == "/auth/login":
            contractor_ok = (
                form.get("contractor") == CONTRACTOR
                if self.server.require_contractor
                else True
            )
            if (
                form.get("userid") == USERNAME
                and form.get("password") == PASSWORD
                and contractor_ok
            ):
                self.send_response(302)
                self.send_header("Set-Cookie", "tegg-session=1; Path=/")
                self.send_header("Location", "/sales/gm-dashboard")
                self.end_headers()
                return
            return self._send(self._login_page(error="Invalid credentials."))

        if not self._authed():
            return self._redirect("/auth/login")

        # Certificate generation.
        if path == "/document-library/certificates":
            if not form.get("solution"):
                return self._send(self._certificates_page(error="Select a Solution."))
            return self._send(
                make_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                {"Content-Disposition": 'attachment; filename="Certificates4471.docx"'},
            )

        # Report: "Print Report" stages the viewer, "Export" delivers the file.
        if path.startswith("/reports/standard-esa/"):
            slug = path.rsplit("/", 1)[-1]
            name = SLUG_TO_REPORT.get(slug)
            if not name:
                return self.send_error(404)

            if form.get("action") == "print":
                missing = [
                    f for f in REPORT_FIELDS[name]
                    if not form.get(f.lower().replace(" ", "-"))
                ]
                if missing:
                    return self._send(
                        self._report_page(
                            name, False, error=f"Required: {', '.join(missing)}"
                        )
                    )
                return self._redirect(f"/reports/standard-esa/{slug}?staged=1")

            if form.get("action") == "export":
                if form.get("select-a-format") != "Acrobat PDF file":
                    return self._send(
                        self._report_page(name, True, error="Choose a format.")
                    )
                _, filename, pages = REPORTS[name]
                return self._send(
                    make_pdf(filename, pages),
                    "application/pdf",
                    {"Content-Disposition": f'attachment; filename="{filename}"'},
                )

        self.send_error(404)

    # -- pages ------------------------------------------------------------
    def _login_page(self, error="") -> bytes:
        banner = f'<p class="error">{error}</p>' if error else ""
        contractor_field = (
            """<p><label for="contractor">Contractor</label>
     <select id="contractor" name="contractor">
       <option value=""></option><option value="Lippolis">Lippolis</option>
       <option value="Other">Other</option></select></p>"""
            if self.server.require_contractor
            else ""
        )
        return _page(
            "Sign In",
            f"""{banner}
<form method="post" action="/auth/login">
  {contractor_field}
  <p><label for="userid">User Id</label>
     <input id="userid" name="userid" type="text"></p>
  <p><label for="password">Password</label>
     <input id="password" name="password" type="password"></p>
  <button type="submit">Log In</button>
</form>""",
        )

    def _documentation_page(self) -> bytes:
        return _page(
            "Documentation",
            f"""<form method="get">
  <p><label for="customer">Search Customer</label>
     <input id="customer" name="customer" type="search"></p>
  {_select("Site Name", SITES)}
</form>
<p><a href="/document-library">Document Library</a></p>""",
        )

    def _certificates_page(self, error="") -> bytes:
        banner = f'<p class="error">{error}</p>' if error else ""
        return _page(
            "Certificates",
            f"""{banner}
<form method="post" action="/document-library/certificates">
  {_select("Solution", AGREEMENTS)}
  <button type="submit">Print/Generate Document</button>
</form>""",
        )

    def _report_page(self, name: str, staged: bool, error="") -> bytes:
        banner = f'<p class="error">{error}</p>' if error else ""
        slug = REPORTS[name][0]
        fields = "".join(
            _select(field, FIELD_OPTIONS[field]) for field in REPORT_FIELDS[name]
        )
        if staged:
            body = f"""{banner}
<form method="post" action="/reports/standard-esa/{slug}">
  <input type="hidden" name="action" value="export">
  {_select("Select a Format", ["Acrobat PDF file", "Excel", "Word"])}
  <button type="submit">Export</button>
</form>"""
        else:
            body = f"""{banner}
<form method="post" action="/reports/standard-esa/{slug}">
  <input type="hidden" name="action" value="print">
  {fields}
  <button type="submit">Print Report</button>
</form>"""
        return _page(name, body)


class MockPortal:
    """Context manager that runs the mock portal on a background thread."""

    def __init__(self, port: int = 0, require_contractor: bool = True) -> None:
        self.server = HTTPServer(("127.0.0.1", port), Handler)
        self.server.require_contractor = require_contractor
        self.port = self.server.server_address[1]
        self.url = f"http://127.0.0.1:{self.port}"

    def __enter__(self) -> "MockPortal":
        self.thread = threading.Thread(target=self.server.serve_forever, daemon=True)
        self.thread.start()
        return self

    def __exit__(self, *exc) -> None:
        self.server.shutdown()
        self.server.server_close()


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--port", type=int, default=8899)
    args = parser.parse_args()
    with MockPortal(args.port) as portal:
        print(f"mock TEGG portal on {portal.url}  (ctrl-c to stop)")
        try:
            threading.Event().wait()
        except KeyboardInterrupt:
            pass
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
