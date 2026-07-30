"""Certificate compatibility.

The real TEGG certificate encodes its checkboxes as Wingdings glyphs, two per
item (Yes and No), across eleven items. The code must recognise that and refuse
to edit it -- an ESA certificate is a legal attestation, and a wrongly ticked
box is worse than an untouched one.

The synthetic fixtures here reproduce that structure exactly, so these tests
run without needing the real customer document on disk.
"""

from __future__ import annotations

import shutil
import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from tegg import certdoc  # noqa: E402
from tegg.certificate import find_soffice  # noqa: E402

HAS_SOFFICE = bool(find_soffice())
REAL_DOC = ROOT / "test-data" / "Certificates71999.doc"

UNCHECKED = certdoc.WINGDINGS_UNCHECKED


def _wingdings_certificate(path: Path, items: int = 11, boxes_per_item: int = 2):
    """A certificate shaped like the real one: glyph checkboxes, Yes/No pairs."""
    from docx import Document

    document = Document()
    document.add_paragraph("CONFIRMATION OF FACILITY ELECTRICAL")
    # The first group, before "A.", which the SOP says to delete.
    document.add_paragraph(f"{UNCHECKED}\tMain entrance switchgear")
    document.add_paragraph(f"under the contract:\t{UNCHECKED}   Yes\t{UNCHECKED}   No")
    document.add_paragraph("A.\tEffective date of service contract:\t______")
    document.add_paragraph("B.\tService contract includes the following:")
    for index in range(1, items + 1):
        boxes = "\t".join(f"{UNCHECKED}   {label}" for label in ("Yes", "No")[:boxes_per_item])
        document.add_paragraph(f"({index})\tItem {index}:\t{boxes}")
    document.save(str(path))
    return path


def _form_field_certificate(path: Path, items: int = 10):
    """A certificate using legacy form fields, one box per item."""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document = Document()
    document.add_paragraph("A.\tEffective date:")
    document.add_paragraph("B.\tIncludes the following:")
    for index in range(1, items + 1):
        paragraph = document.add_paragraph(f"({index})\tItem {index}:")
        run = paragraph.add_run()
        field = OxmlElement("w:fldChar")
        field.set(qn("w:fldCharType"), "begin")
        data = OxmlElement("w:ffData")
        checkbox = OxmlElement("w:checkBox")
        default = OxmlElement("w:default")
        default.set(qn("w:val"), "0")
        checkbox.append(default)
        data.append(checkbox)
        field.append(data)
        run._r.append(field)
    document.save(str(path))
    return path


class TestAnalyzeWingdings:
    @pytest.fixture
    def certificate(self, tmp_path):
        return _wingdings_certificate(tmp_path / "cert.docx")

    def test_detects_glyph_checkboxes(self, certificate):
        analysis = certdoc.analyze(certificate)
        assert certdoc.WINGDINGS_GLYPH in analysis.encodings

    def test_finds_sections_a_and_b(self, certificate):
        analysis = certdoc.analyze(certificate)
        assert analysis.section_a_index is not None
        assert analysis.section_b_index is not None
        assert analysis.section_b_index > analysis.section_a_index

    def test_counts_all_eleven_items(self, certificate):
        assert len(certdoc.analyze(certificate).item_lines) == 11

    def test_sees_two_boxes_per_item(self, certificate):
        assert certdoc.analyze(certificate).boxes_per_item == 2

    def test_refuses_to_call_it_editable(self, certificate):
        analysis = certdoc.analyze(certificate)
        assert analysis.editable is False
        assert "wingdings" in analysis.reason().lower()

    def test_describe_is_readable(self, certificate):
        described = certdoc.analyze(certificate).describe()
        assert "safe to edit" in described
        assert "NO" in described


class TestAnalyzeOtherEncodings:
    def test_form_fields_with_one_box_per_item_are_editable(self, tmp_path):
        certificate = _form_field_certificate(tmp_path / "ff.docx")
        analysis = certdoc.analyze(certificate)
        assert certdoc.LEGACY_FORM_FIELD in analysis.encodings
        assert analysis.editable is True

    def test_a_yes_no_pair_layout_is_never_editable(self, tmp_path):
        """Even a supported encoding is unsafe if numbering cannot be mapped."""
        certificate = _wingdings_certificate(tmp_path / "pairs.docx")
        analysis = certdoc.analyze(certificate)
        analysis.encodings = {certdoc.LEGACY_FORM_FIELD}
        assert analysis.boxes_per_item == 2
        assert analysis.editable is False

    def test_a_document_with_no_checkboxes_is_not_editable(self, tmp_path):
        from docx import Document

        path = tmp_path / "plain.docx"
        document = Document()
        document.add_paragraph("A.\tnothing here")
        document.save(str(path))
        analysis = certdoc.analyze(path)
        assert analysis.editable is False
        assert "no checkboxes" in analysis.reason()

    def test_a_legacy_doc_is_rejected_with_a_message_not_a_traceback(self, tmp_path):
        path = tmp_path / "legacy.doc"
        path.write_bytes(b"\xd0\xcf\x11\xe0legacy binary")
        analysis = certdoc.analyze(path)
        assert analysis.error
        assert "not a .docx" in analysis.error
        assert analysis.editable is False

    def test_an_unopenable_docx_reports_an_error(self, tmp_path):
        path = tmp_path / "broken.docx"
        path.write_bytes(b"PK\x03\x04 not really a docx")
        analysis = certdoc.analyze(path)
        assert analysis.error
        assert analysis.editable is False


@pytest.mark.skipif(not HAS_SOFFICE, reason="LibreOffice not installed")
class TestConversion:
    def test_docx_is_returned_unchanged(self, tmp_path):
        certificate = _wingdings_certificate(tmp_path / "cert.docx")
        assert certdoc.ensure_docx(certificate, tmp_path) == certificate

    def test_prepare_produces_a_pdf_and_demands_review(self, tmp_path):
        certificate = _wingdings_certificate(tmp_path / "cert.docx")
        work = tmp_path / "converted"
        outcome = certdoc.prepare(certificate, work)
        assert outcome.pdf.exists()
        assert outcome.pdf.name == "Certificates good.pdf"
        assert outcome.edited is False
        assert outcome.review_required is True
        assert "by hand" in outcome.reason

    def test_prepare_never_modifies_the_original(self, tmp_path):
        certificate = _wingdings_certificate(tmp_path / "cert.docx")
        before = certificate.read_bytes()
        certdoc.prepare(certificate, tmp_path / "converted")
        assert certificate.read_bytes() == before

    def test_missing_source_is_a_clear_error(self, tmp_path):
        from tegg.certificate import CertificateError

        with pytest.raises(CertificateError, match="not found"):
            certdoc.convert(tmp_path / "nope.docx", tmp_path)


@pytest.mark.skipif(not REAL_DOC.exists(), reason="no real certificate available")
@pytest.mark.skipif(not HAS_SOFFICE, reason="LibreOffice not installed")
class TestRealCertificate:
    """Against the actual document downloaded from the portal."""

    def test_legacy_doc_converts_and_is_classified_unsafe(self, tmp_path):
        source = tmp_path / REAL_DOC.name
        shutil.copy2(REAL_DOC, source)
        docx = certdoc.ensure_docx(source, tmp_path)
        assert docx.suffix == ".docx"

        analysis = certdoc.analyze(docx)
        assert certdoc.WINGDINGS_GLYPH in analysis.encodings
        assert analysis.boxes_per_item == 2
        assert len(analysis.item_lines) == 11
        assert analysis.editable is False
