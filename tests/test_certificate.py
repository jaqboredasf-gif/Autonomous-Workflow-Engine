import sys
from datetime import date
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))
sys.path.insert(0, str(ROOT / "tests"))

from docx import Document  # noqa: E402
from mock_portal import make_docx  # noqa: E402

from tegg.certificate import (  # noqa: E402
    CertificateError,
    edit_certificate,
    inspect,
)

YES = [1, 2, 3, 4, 5, 9]
NO = [6, 7, 8, 10]


@pytest.fixture
def certificate(tmp_path):
    path = tmp_path / "Certificates4471.docx"
    path.write_bytes(make_docx())
    return path


class TestInspect:
    def test_finds_every_checkbox(self, certificate):
        structure = inspect(certificate)
        # 3 preliminary + 10 section B
        assert len(structure.boxes) == 13

    def test_locates_sections_a_and_b(self, certificate):
        structure = inspect(certificate)
        assert structure.section_a_index is not None
        assert structure.section_b_index is not None
        assert structure.section_a_index < structure.section_b_index

    def test_splits_the_first_group_from_section_b(self, certificate):
        structure = inspect(certificate)
        assert len(structure.first_group) == 3
        assert len(structure.section_b_boxes) == 10

    def test_describe_is_readable(self, certificate):
        text = inspect(certificate).describe()
        assert "first group (delete)  : 3" in text
        assert "section B items       : 10" in text

    def test_missing_file(self, tmp_path):
        with pytest.raises(CertificateError, match="not found"):
            inspect(tmp_path / "nope.docx")


class TestEdit:
    def test_applies_the_sop_answers(self, certificate, tmp_path):
        out = tmp_path / "edited.docx"
        edit = edit_certificate(
            certificate, out, yes_items=YES, no_items=NO, when=date(2026, 7, 29)
        )
        assert edit.boxes_checked == sorted(YES)
        assert edit.boxes_unchecked == sorted(NO)
        assert edit.date_written == "07/29/2026"
        assert out.exists()

    def test_checkbox_states_are_written(self, certificate, tmp_path):
        from docx.oxml.ns import qn

        out = tmp_path / "edited.docx"
        edit_certificate(certificate, out, yes_items=YES, no_items=NO)

        structure = inspect(out)
        assert len(structure.section_b_boxes) == 10

        def state(box):
            for tag in ("w:checked", "w:default"):
                node = box.element.find(qn(tag))
                if node is not None:
                    return node.get(qn("w:val"))
            return None

        for item in YES:
            assert state(structure.section_b_boxes[item - 1]) == "1", f"item {item}"
        for item in NO:
            assert state(structure.section_b_boxes[item - 1]) == "0", f"item {item}"

    def test_first_group_is_removed(self, certificate, tmp_path):
        out = tmp_path / "edited.docx"
        edit = edit_certificate(certificate, out, yes_items=YES, no_items=NO)
        assert edit.boxes_deleted == 3
        assert len(inspect(out).boxes) == 10

    def test_first_group_can_be_kept(self, certificate, tmp_path):
        out = tmp_path / "edited.docx"
        edit = edit_certificate(
            certificate, out, yes_items=YES, no_items=NO, delete_first_group=False
        )
        assert edit.boxes_deleted == 0
        assert len(inspect(out).boxes) == 13

    def test_date_lands_in_section_a(self, certificate, tmp_path):
        out = tmp_path / "edited.docx"
        edit_certificate(
            certificate, out, yes_items=YES, no_items=NO, when=date(2026, 7, 29)
        )
        paragraphs = [p.text for p in Document(str(out)).paragraphs]
        section_a = next(p for p in paragraphs if p.strip().startswith("A."))
        assert "07/29/2026" in section_a

    def test_pinned_group_size_mismatch_is_refused(self, certificate, tmp_path):
        """An explicit size that disagrees with the document must not proceed."""
        with pytest.raises(CertificateError, match="expected a first group of 5"):
            edit_certificate(
                certificate, tmp_path / "o.docx",
                yes_items=YES, no_items=NO, first_group_size=5,
            )

    def test_pinned_group_size_matching_is_accepted(self, certificate, tmp_path):
        edit = edit_certificate(
            certificate, tmp_path / "o.docx",
            yes_items=YES, no_items=NO, first_group_size=3,
        )
        assert edit.boxes_deleted == 3

    def test_refuses_when_section_b_is_too_short(self, tmp_path):
        from docx import Document as NewDoc

        document = NewDoc()
        document.add_paragraph("A.  Date:")
        document.add_paragraph("B.  Items")
        path = tmp_path / "thin.docx"
        document.save(str(path))
        with pytest.raises(CertificateError, match="no form checkboxes"):
            edit_certificate(path, tmp_path / "o.docx", yes_items=YES, no_items=NO)
